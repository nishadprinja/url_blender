import sqlite3
import validators
import requests
import docx
import re
import spacy
import urllib.request
import io
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from validators import ValidationError

def is_string_a_url(url_string):
    result = validators.url(url_string)

    if isinstance(result, ValidationError):
        return False

    return result

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

# Create a SQL connection to our SQLite database
con = sqlite3.connect("/Users/nishadprinja/Library/Messages/chat.db")

cur = con.cursor()
rows = cur.execute('SELECT datetime (message.date / 1000000000 + strftime ("%s", "2001-01-01"), "unixepoch", "localtime") AS message_date, message.text, message.is_from_me, chat.chat_identifier FROM chat JOIN chat_message_join ON chat. "ROWID" = chat_message_join.chat_id JOIN message ON chat_message_join.message_id = message. "ROWID" WHERE chat_identifier = "+18457097580" AND is_from_me = "0"')

doc1 = Document()
doc2 = Document()
doc3 = Document()

list_paragraph1 = doc1.add_paragraph()
list_paragraph2 = doc2.add_paragraph()
list_paragraph3 = doc3.add_paragraph()

rowz = rows.fetchmany(30)

# The result of a "cursor.execute" can be iterated over by row

for i, row in enumerate(rowz):
    if is_string_a_url(row[1]):

        try:
            # Step 1: Fetch the web page
            url = row[1]
            response = requests.get(url)

            if response.status_code == 200:
                # Step 2: Parse the HTML content
                soup = BeautifulSoup(response.text, 'html.parser')

                # Step 3: Extract metadata
                title = soup.find("meta", property="og:title")
                image = soup.find("meta", property="og:image")
                keywords = [item['content'] for item in soup.select('[name=Keywords][content], [name=keywords][content]')]

                title_content = title["content"] if title else "No meta title given"
                image_content = image["content"] if image else "No meta image given"

                """
                print(url)
                print(title_content)
                print(image_content)
                print("\n")
                """

                response = requests.get(image_content, stream=True)
                image_parsed = io.BytesIO(response.content)
                
                print(image_parsed)

                general_keywords = ['video', 'sharing', 'camera phone', 'video phone', 'free', 'upload']

                entertainment_keywords = ['kurt angle', 'kurt angle interview', 'wwe shorts', 'kurt angle milk truck', 'kurt angle shoot interview', 'kurt angle milk chug', 'Kurt Angle Milk story', 'Kurt Angle Funny Story', 'shoot interview', 'kurt angle wwe hall of fame', 'kurt angle funny', 'wrestling shorts', 
                'jon stewart', 'jon stewart stephen colbert', 'jon stewart on stephen colbert', 'stephen colbert', 'stephen colbert show', 'stephen colbert jon stewart', 'colbert jon stewart interview', 'jon stewart interview', 'jon stewart covid', 'jon stewart coronavirus', 'jon stewart covid origin', 'wuhan', 'wuhan china', 'coronavirus', 'wuhan coronavirus', 'jon stewart covid lab reaction', 'jon stewart react', 'colbert show', 'jon stewart colbert show', 'covid origin solved', 'shorts', 'yaf', 'young americas foundation', 
                'Johnny Depp', 'live', 'stage', 'jeff beck', 'SNL', 'Saturday Night Live', 'sketches', 'funny', 'comedy', 'only', 'musical guest', 'host', 'saturday', 'live', 'new york', 'comedian', 'snline', 'trivia', 'fans', 'waiting in line', 'humor', 'sketch', 'sketch comedy', 'Dwayne Johnson', 'The Rock', 'George Ezra', 'bobby moynihan', 'WWE', 'Wrestling (Sport)', 'promo shoot', 
                'dave grohl interview', 'dave grohl drumming', 'foo fighters', 'dave grohl', 'nirvana', 'foo fighter everlong', 'foo fighter best songs', 'foo fighter playlist', 'nirvana songs', 'best of nirvana', 'best of foo fighters', 
                'funny', 'meme', 'memes', '4chan', 'shitpost', 'based', 'SM64', 'sm64 soundfont', 'super mario 64', 'n64', 'nintendo 64', 'nintendo', 'alice in chains', 'aic', 'dirt', 'facelift', 'sap', 'jerry cantrell', 'nirvana', 'pearl jam', 'soundgarden', 'grunge', 'music', '90s music', 
                'heavy metal', 'death metal', 'grunge', 'grunge band', 'nirvana', 'alice in chains', 'aic', 'mike starr', 'layne staley', 'jerry cantrell', 'sean kinney', 'good vibes', 'good', 'concert', 'live', 'shorts', 'shorts music', 'shorts video', 'alternative music', 'rock', 'rock metal', 'rock band', 'rock vibes', 'rock music band', 'band', 'live music', 'good ending', '90s', '80s', 'music 90s', 'music 80s', '80s music', '90s music', 'usa', 'international', 
                'drums', 'drummers', 'drummer', 'how drummers knock on doors', 'funny', 'jokes', 'comedy', 'music', 'drum beats', 'name this song', 'name that song', 'name that beat', 'what song is this', 'metallica', 'disturbed', 'down with the sickness', 'metallica one', 'lars ulrich', 'led zeppelin', 'eye of the tiger', 'lorna shore', 'slipknot', 'joey jordison', 'meshuggah bleed', 'bleed drum cover', 'taiko no tatsujin', 'i hear your heartbeat', 'drums fun', 'drum cover', 
                'Toad', 'Memes', 'Micheal jackson', 'Pain', 'Billie jean', 'Music', 'Meme', 'Impressions', 'Impression', 'Mario', 'Super Mario', 'Mario64', 'Cursed', 'N64', 'Ass', 'Joke', 'Trash', 'Garbage', 'Singing', 'Singer', 'Musician', 'Pop', 'Rock', 'Mushroom?', 'Cover', 
                'bernth drown', 'amp under water', 'guitar amp under water', 'guitar filled with water', 'bernth waterworks', 'bernth album', 'bernth band', 
                'Breaking Bad', 'Walter White', 'Bryan Cranston', 'Anna Gunn', 'Aaron Paul', 'Dean Norris', 'Betsy Brandt', 'Bob Odenkirk', 'Jonathan Banks', 'Jesse Pinkman', 'Official Trailer', 'Vince Gilligan', 'Breaking Bad full episodes', 'Breaking Bad clips', 'Breaking Bad funny moments', 'Breaking Bad bloopers', 'Better Call Saul', 'Breaking Bad Season', 'how to spot a counterfeit bill ft. nacho varga better call saul', 'how to spot a counterfeit bill 100', 'nacho varga', 'better call saul behind the scenes', 
                'mothership', 'john cena', 'wwe', 'champion', 'love', 'lao gan ma', 'China', 'speaks', 
                'guitar', 'arrangement', 'fingerstyle', 'fingerstyle guitar', 'tik cover', 'guitar cover', 'guitarist', 'funny', 'funny guitar', 'acoustic guitar', 'tabs', 'guitar music', 'Song', 'When', 'when you', "When you don't know a single chord", 'but you need to impress new friends', 'to impress', 
                'daft punk', 'i feel it coming', 'the weeknd', 'weeknd', 'pee is stored in the balls', 'pee', 'stored', 'balls', 'memes', 'grammys', 'discovery', 'get lucky', 'starboy', 
                'Dorkly', 'Video Game Culture', 'videogames', 'sonic', 'dorkly bits', 'huh neat', 'sprites', 'super punch out', 'mike tyson', 'mario', 'link', 'zelda', 'ganon', 'bowser', 'smart', "voldemort's assistant kevin", 'mega man', 'space invaders', 'humor', 'nintendo', 
                'advert', 'atari', 'retro', 'jack black', 'wallpapers', 'vintage', 'tv', '2600', 'pitfall', 'funny', 
                'Inside The Ropes', 'Kenny McIntosh', 'bret hart', 'kevin nash', 'montreal screwjob', 'the kliq', 'shawn michaels', 'razor ramon', 'diesel', 'scott hall', 'sean waltman', '123 kid', 'x-pac', 'bret hitman hart', 'hitman', 'Twitter', 'Montreal', 'Screwjob', 'screw', 'job', 'Bret Hart Shorts', 'WWE Shorts', 'Wrestling Shorts', 'Bret Hart Shoot Interview', 'Shoot', 'Interview', 'Wrestling', 'WWE', 'HBK', 'Earl Hebner', 
                'Super Smash Bros', 'Super Smash Bros Ultimate', 'Super Smash Bros Meme', 'Super Smash Bros Ultimate Meme', 'Country Roads Mario', 'nintendo meme', 
                'Brock Lesnar', 'Kurt Angle', 'WWE', 'Milk', 'instagram', 'kingdom hearts', 'sora', 'riku', 'cutscene', 'english', 'cutscenes', 'dearly beloved', 'traverse town', 'music', 'darkness', 'dark', 
                'bret', 'hart', 'best', 'athletic', 'moves', 'tribute', 'the', 'hit', 'man', 'hitman', 'athlete', 'athleticism', 'speed', 'agility', 'high', 'risk', 'taking', 'top', 'maneuvers', 'wwf', 'wwe', 'stampede', 'wrestling', 'ecw', 'pro', 'professional', 'sports', 'entertainer', 'entertainment', 'aerial', 'japan', 
                'The Howard Stern Show', 'Howard Stern', 'Robin Quivers', 'Fred Norris', 'Baba Booey', 'SiriusXM', 'HSS', 'htvod', 'howard stern show', 'howard stern interviews', 'stern show', 'living colour', 'living color', 'cult of personality', 'rock band', 'rock music', 
                'America\'s Got Talent', 'America\'s Got Talent The Champions', 'Season 2', 'Talent', 'The Champions', 'NBC', 'Got Talent', 'Singing', 'Auditions', 'Live Performance', 'Full Performance', 'America\'s', 'American Idol', 'Singer', 'Audition', 'Abc', 'Season 15', 'Agt', 'Howie Mandel', 'Simon Cowell', 'Sofia Vergara', 'Heidi Klum', 'Terry Crews', 'Golden', 'Buzzer', 
                'family', 'guy', 'theme', 'song', 'alive', 'live', 'inside', 'actors', 'studio', 'seth', 'mcfarlene', 'alex', 'borstein', 
                'system', 'of', 'down', 'soad', 'system of a down', 'daron', 'malakian', 'serj', 'tankian', 'john', 'dolmayan', 'shavo', 'odadjian', 'scars', 'on', 'broadway', 'north', 'kingsley', 'these', 'grey', 'men', 'armenia', 'yerevan', 'music', 'live', 'performance', 'rare', 'stage', 'song', 'new', 'scars on broadway', 'north kingsley', 'these grey men', 'full', 'hd', 'velvet', 'hammer', 'california', 'los', 'angeles', 'la', 'show', 'concert', 'interview', 'chop', 'suey', 'chop suey!', 'isolated', 'vocals', 
                'rock', 'guitar', 'drummer', 'reupload', 'bass', 'drum cover', 'drumming', 'drums', 'percussion', 'drum solo', 'Metal', 'music', 'metal', 'djent', 
                'Comedy', 'Humor', 'Shorts', 'YouTube shorts', 'Viral', 'Family friendly', 'notorious big', 'big', 'biggie smalls', 'p diddy', 'diddy', 'charlie bahama', 'biggie', 'puff daddy', 'hip hop', '90s hip hop', 'rap music', 'electric air', 'earthbeat films', 'the bahamas', 'the charlie bahama show', 'iconic hip hop', 'classic rap music', 'icons', 'bad boy records', 'michael jackson', 'puffy', 'jackson five', 'jackson 5', 'this is it', 'the king of pop', 'junior maffia', 'bad boy', 
                'linkin park', 'linkin park hits', 'linkin park albums', 'linkin park songs', 'chester bennington', 'meteora', 'meteora linkin park full album', 'meteora linkin park', '8 bit', 'chiptune', '8 bit edit', 'remix', 'edit', 'linkin park meteora', 'album 8 bit', 'dont stay', 'numb', 'somewhere i belong', 'faint', 'easier to run', 'breaking the habit', 'album', 
                'Best Doritos Commercial Ever', 'hilarious', 'funny', 'Comedy', 'Doritos', 'laugh', 'Super Bowl', 'Knee Slapper', 'commercials', 'Spot', 'Ads', 'commercial', 'Best', 'Cinema 7 Pictures', 'Austin Ellsworth', 'Super', 'Bowl', 'Beer', 'Nacho Cheese', 'Nacho', 'Cheese', 'Sexy', 'Explosion', 'Super Bowl', 'competition', 'winner', 'superbowl', 'hot', 'comedy', 'boobs', 'winning', 'crash the superbowl', 'banned', 'crash the super bowl', 'cheesemaker', 'girl', 
                'crispy concords', 'crispy', 'shorts', 'extra crispy', 'omegle', 'make me laugh', 'joke', 'prank', 'interview', 'social experiment', 'the undertaker', 'undertaker', 'wwe', 'undertaker hates', 'wwe taker', 'wwe undertaker', 'undertaker list', 'wwe list', 'hulk hogan', 'sting', 'shawn michaels', 'undertaker hate', 'triple h', 'cm punk', 'cm punk aew', 'cm punk fired', 'Dog', 'Burrito', 'Bean', 'Taco', 'Bell', 'Funny', 'Gross', 
                'comedy', 'comedy shorts', 'funny moments', 'videos to laugh', 'funny shorts', 'comedy club', 'stand-up', 'Dad', 'Dan', 'father', 'son', 'movie', 'cinematic', 'dads', 'sons', 'Michael', 'Michael Jackson', '80s', 'icons', 'MTV', 'we are the world', 'backstage', 'recording', 'new york city', 'best coffee in nyc', 'new york coffee shops', 'new york', 'caffeine test', 'caffeine levels', 'test caffeine', 'coffee shops', 'caffeine mg', 'cold brew', 'coffee', 'caffeine', 'coffee in new york', 'nyc vlog', 
                'coffee in nyc', 'new york cafe', 'nyc coffee', 'espresso', 'caffeine dosage', 'fda', 'james hoffmann', 'artizan coffee', 'stardust coffee', 'astrocat app', 'caffeine tracker app', 'caffeine app', 'coffee app', 'Lighttells CA-700 Caffeine', 'lighttells caffeine analyzer', 'Ryan Stiles', 'Colin Mochrie', 'best of ryan stiles', 'ryan and colin', 'ryan stiles colin mochrie', 'best of colin mochrie', 'whose line is it anyway', 'funniest ryan stiles moments', 'whose line ryan', 'best of whose line', 'ryan stiles top 10', 'top 10', 
                'world wrestling entertainment', 'wwe', 'paul heyman', 'aew', 'wrestling', 'wwe interview', 'wrestle', 'wrestler', 'superstar', 'shorts', 'short', 'short video', 'tik tok', 'triple h', 'brock lesnar', 'seth rollins', 'rollins', 'ROMAN REIGNS', 'nikki bella', 'the miz', 'roman reigns', 'smackdown', 'cena', 'the shield', 'sasha banks', 'This Was NOT Supposed To Happen', 'cody rhodes', 'brock cody', 'cody rhodes brock lesnar', 
                'n64', 'nintendo 64', 'n64 mods', 'n64 hacks', 'Eminem', 'Slim shady', 'marshall mathers', 'Rap', 'parody', 'skit', 'tiktok', 'youtube shorts', 'viral', 'true story', 'parents', 'parents when', 'parents memes', 'memes', 'dank', 'horror', 'haunted house', 'haunted', 'Music cover', 'official', 'music video', 'Piano', 'Pianist', 'Concert', 'Singing', 'Singer', 'Cover', 'Performance', 'American Girl', 'Girl of the year', 'vlog', 'video blog', 'vlogging', 'vlogger', 'daily vlog', 'daily', 'family', 'baby', 'funny video', 'funny', 'funniest', 'funny babies', 'funny cats', 
                'pregnancy', 'pregnant', 'birth', 'review', 'ootd', 'beauty', 'outfit', 'curly', 'style', 'disney', 'how to', 'Justin Bieber', 'Justin', 'Bieber', 'Love', 'Yourself', 'Love Yourself', 'WWE Wrestlers Who Suffered With Illnesses & Disorders', 'Wrestlers Who Suffered With Illnesses & Disorders', 'WWE Wrestlers', 'WWE', 'wwe', 'wwe wrestlers who died', 'wwe night of champions 2023', 'wwe night of champions 2023 match card', 'wwe night of champions 2023 predictions', 'wwe night of champions 2023 winner prediction', 'wwe news', 'wwe news and rumors', 'wwe rumors', 'pro wrestling rumors', 
                'funny cat videos', 'cat behavior', 'cat antics', 'amusing cats', 'cat reactions', 'cute cat videos', 'entertaining pets', 'viral pet videos', 'cat and empty bowl', 'humorous pet moments', 'funny animal videos', 'amusing pet videos', 'adorable cats', 'pet humor', 'pet and owner moments', 'entertaining animal videos', 'cat and small spaces', 'cat and bowls', 'cat and fitting inside', 'cat purring', 'if I fits', 'I sits', 'pet and comfort', 'pet instinct', 'heartwarming pet moments', 'cats', 'if i fits i sits', 'animals', 'kittens', 
                'Cheater', 'Caught cheating', 'Cheating', 'Relationship', 'Couple', 'Married', 'Marriage', 'Sad', 'england', 'rehearsal', 'open', 'girl', 'ohio', 'princess', 'front', 'door', 'heartwarming', 'prom', 'teens', 'creative', 'expecting', 'medieval', 'messenger', 'promposal', 'text', 'decided', 'latest videos', 'viral videos', 'viral', 'fail videos', 'crashes', 'stunts', 'amazing', 'animal attack', 'world record', 'news flare', 'high school', 'biker', 'bikers', 'motorcycle', 'motorcycles', 'motovlogs', 'sportbikes', 'sport bikes', 'bike life', 'arizona', 'bikers are cool', 'bikers are awesome', 'bikers are nice', 'biker videos', 'sport bike videos', 'riding a motorcycle', 'sportbike videos', 'biker shorts', 'gopro', 
                'Daily Dose Of Internet', 'worlds cheapest', "The World's Cheapest Millionaire", 'worlds cheapest people', 'worlds cheapest millionaire', 'worlds cheapest woman', 'the worlds cheapest', 'cheapest woman ever', 'cheapest woman', 'cheapest millionaire', 'cheapest people ever', 'aimee elizabeth millionaire', 'aimee elizabeth extreme cheapskates', 'the cheapest millionaire', 'the cheapest woman in the world', 'stand up comedy', 'stand up', 'dont tell comedy', 'standup comedy', 'stand up comedian', 'don\'t tell comedy', 'best stand up', 'comedy specials', 'best stand up comedy', 'jokes', 'funny comedy specials', 'best comedy specials', 'comedy videos', 'stand up comedy full', 'stand up comedy full special', 
                'animated', 'animation', 'animation movies', 'batman', 'batman animated', 'christopher nolan', 'kevin conroy', 'talkin toons', 'the dark knight', 'the dark knight animated', 'Netflix', 'Australia', 'archive', 'archival footage', 'retrofocus', 'abc news', 'abc tv', 'from the archive', 'looks', 'attraction', 'dating', 'ugly', 'ugly men', 'vox pops', 'melbourne', 'melbourne history', 'Taylor Swift', 'LGBT', 'rock', 'guitar', 'drummer', 'bass', 'drum cover', 'drumming', 'drums', 'percussion', 'drum solo', 'Metal', 'music', 'metal', 'djent', 
                'youtube poop', 'intro meme', 'meme', 'memes', 'dank memes', 'meme compilation', 'rickroll', 'dbz', 'dbz meme', 'goku', 'goku meme', 'kakarot', 'vegeta', 'kakarot meme', 'dragonball', 'dragonball z', 'dragonball z meme', 'vegeta meme', 'frieza meme', 'Superstars', 'raw', 'smackdown', 'afterburn', 'nxt', 'ppv', 'pay per view', 'divas', 'Classic Rock', 'microphone mayhem', 'the Great One', 'Stone Cold Steve Austin', 'Undertaker', 'Hell in a Cell', 'Monday Night Raw', 'spiderman', 'tobey', 'maguire', 'bully', 'mcu', 'mighty morphin power rangers', 'power rangers', 'power rangers movie', '90s TV shows', '90s', 'nostalgia', 'guitar', 'guitar cover', 
                'knuckles', 'meme', 'vrchat', 'the way', 'male fantasy', 'video games', 'gaming', 'WCW/nWo', 'WCW', 'nWo', 'WCW/nWo Revenge', 'WCW Revenge', 'WCW nWo Revenge', 'WCW nWo', 'Revenge with Music', 'WCW Revenge Theme Songs', 'WCW Music', 'aew', 'aew dynamite', 'aew dynamite highlights this week', 'aew dynamite highlights', 'aew on tnt', 'aew road to', 'all elite wrestling', 'pro wrestling', 'aew on tbs', 'aew rampage', 'aew highlights', 'aew rampage highlights', 'aew zero hour', 'aew buy in', 'Adam', 'Copeland', 'Rated R', 'Superstar', 'Edge', 'Christian', 'Cage', 'Christian Cage', 'Adam Copeland', 'Rated R Superstar', 'Metalingus', 'sting', 'Darby', 'Allin', 
                'custom ps5', 'ps5 cheap custom', 'cheap ps5', 'led ps5', 'stickers ps5', 'skin ps5', 'ps5 pro', 'ps5 unboxing', 'ps5 slim', 'psp', 'psp ps5', 'ps remote play', 'learn about wolfdogs', 'Bob and the wolf dogs', 'wolfdogs', 'wolf hybrid', 'pet wolf', 'wolf dog hybrid', 'snow dogs', 'wolf', 'wolves', 'wolf Honey', 'wolf Puppy', 'animal video', 'animals', 'the dodo', 'Animal Rescue', 'dodo', 'cute animals', 'pets', 'wildlife', 'pet videos', 'wildlife videos', 'animals the dodo', 'the dodo animals', 'rescuing animals', 'ps5', 'playstation 5', 'playstation', 
                'retro gaming', 'gameboy advance', 'the retro future', 'boxy pixel', 'retro gaming now', 'retro gaming console', 'retro gaming setup', 'game boy', 'gameboy mods', 'gameboy mods shop', 'apple', 'iPhone', 'new iphone', 'smartphone', 'Whose Line is it Anyway', 'funny', 'ryan stiles', 'colin mochrie', 'drew carey', 'wayne brady', 'Picking up girls', 'Pickup lines', 'Family friendly', 'All Your Pain in One Video', 'All', 'Pain', 'All Your Pain', 'in One Video', 'funny', 'funniest', 'real life', 'life', 'all your pain in video', 'all your pain in 1 video', 'all pain in one video', 'shot on iPhone meme', 'WHEN IT HURTS', 
                'headbang', 'head bang', 'longhair', 'long hair', 'Star Wars', 'The Last Jedi', 'The Rise of Skywalker', 'The Force Awakens', 'Kylo Ren', 'Luke Skywalker', 'Rey Palpatine', 'Finn', 'Poe Dameron', 'Leia Organa', 'Darth Vader', 'Baby Yoda', 'The Mandalorian', 'Chewbacca', 'party', 'jre', 'jre clips', 'joe rogan', 'joe rogan experience', 'joe rogan podcast', 'ufc', 'pride fc', 'bellator', 'dana white', 'don frye', 'chael sonnen', 'brock lesnar', 'combat', 'sports', 'wwf', 'wcw', 'wrestling', 'kickboxing', 'boxing', 'ufc history', 'ufc legend', 'mma legend', 'mma history', 'dana white ufc', 'ufc interview', 'mma interview', 'ufc podcast', 'mma podcast', 'conor', 'mcgregor', 'notorious', 'ufc fight', 'ultimate fighting championship', 'mma clips', 'mma', 'ufc', 'frank mir', 
                'dave grohl', 'nirvana', 'scream', 'foo fighters', 'band', 'punk', 'rock', 'guantanamo bay', 'Harold and Kumar', 'BEST MOVIE SCENE EVER', 'airplane', 'scene', 'harold', 'RKO', 'Randy Orton', 'Bob Orton', 'Roman Reigns', 'CM Punk', 'Kevin Nash', 'DDP', 'Vince Russo', 'Dave Meltzer', 'Maven', 'John Cena', 'Cowboy Bob Orton', 'Wrestling', 'Diamond Cutter', 'Ace Crusher', 'Randy Orton Top 10', 'The Best', 'ECW', 'AEW', 'Wrestlemania', 'Survivor Series', 'Royal Rumble', 'Kenny Omega', 'MJF', 'highlights', 'RARE', 'Don Muraco', 'Triple H', 'Bianca Belair', 'Becky Lynch', 'Rey Mysterio', 'Bret Hart', 'Shoots on', 'Alexa Bliss', 'The Viper', 'Legend Killer', 'Bloodline', 'NJPW', 'Okada', 'Top 10', 
                'Chester Bennington', 'Linkin Park', 'Chester of Linkin Park', 'Young Chester Bennington', 'Pranks', 'entertainment', 'top 20 guitar riffs', 'top 20 guitar riffs of all time', 'top 50 guitar riffs', 'four horsemen', 'dusty rhodes', 'ultimate warrior', 'macho man', 'scott steiner', 'steiners', 'steiner brothers', 'hulkamania', 'ric flair', 'The Legend of Zelda', 'Nintendo', 'Ocarina of Time', 'Shrek', 'rocksmith', 'rocksmith+', 'review', 'instruments', 'National Football League', 'NFL', 'NFC West', 'Football', 'NFL News', 'Trump', 'Donald Trump', 'potus', 'Dorkly', 'lol', 'pokemon', 'dorkly pokemon', 'animation', 'literal pokedex', 'dorkly bits pokemon', 'dorkly bits pokedex', 'literal pokedex entries', 'pokedex entries', 'pokemon funny', 'pokemon parody', 
                'pokedex animated', 'marvel', 'mcu', 'eternals', 'antman', 'spider-man', 'marvel cinematc universe', 'captain', 'captain marvel', 'the marvels', 'marvels', 'ant-man', 'ant man', 'Kurt cobain', 'Nirvana', 'Dave grohl', '10 hours', 'tenacious d', 'tribute', 'kyle gass', 'KG', 'King Jables', 'JB', 'Jables', 'Jablinski', 'Jablinski Games', 'IPad', 'Halloween Costume', 'fake blood', 'Halloween', 'iPad2', 'Facetime', 'videochat', 'see through', 'illusion', 'simple', 'mobile mifi', 'duct tape', 'torso', 'costume', 'bloody', 'gory', 'Mark Rober', 'NASA', 'engineer', 'Mercedes', 'Fuel Cell', 'Electric Car', 'Publicity Stunt', 'Invisible', 'Video trick', 
                'Internet', 'Internet videos', 'Daily Dose', 'TikTok Compilation', 'Compilation', 'music', 'new music', 'independent artist', 'emo', 'midwest emo', 'punk rock', 'nostalgia', 'cover song', 'jack sparrow', 'disney characters', 'Peekaboo Cat Cave', 'Cat Cave Adventure', 'Feline Hideout', 'Cat Playtime', 'Interactive Cat Toys', 'Pet Accessories', 'Cat Lover\'s Paradise', 'Cozy Cat Bed', 'Cute Kittens', 'Playful Pets', 'Cat Tunnel', 'Cat Store', 'Cattasaurus', 'Cat Furniture', 'Cat Hideout', 'Cat Products', 'Cat facts', 'Cat Tips', 'Feline Tips', 'Heathy Cat', 'Bored Cat', 'Feline Comfort', 'Kitty Retreat', 'Cat Enrichment', 'Innovative Cat Cave', 'Pet Furniture', 'Indoor Cat Cave', 'Cat House', 'Peekaboo Cat Bed', 'Felt Cat Cave', 'Cat Hideaway', 
                'parrot', 'bird', 'quirky', 'bird ukulele', 'parrot sings', 'pokemon', 'pokemon cartridge', 'pokemon event distribution cartridge', 'pokemon games', 'pokemon events', 'celebrity', 'pizza', 'pizza vending machine', 'japan', 'japanese', 'asia', 'vending machine', 'singapore', 'fresh pizza', 'handmade pizza', 'Vending machine', 'challenge', 'PS2', 'PS2 sound', 'ps2 meme', 'video game meme', 'Japanese food', 'japan travel', 'street food', 'Japanese life', 'mario', '64', 'sm64', 'game', 'nintendo', 'nintendo 64', 'impossible 1up', 'collecting the impossible coin', 'mario online', 'super mario', 'super mario bros', 'super mario movie', 'mario challenge', 'funny mario videos', 'mario 64', 'Mario 64 trick', 'mario 64 speedrun', 'mario 64 speedrun trick', 'mario 64 tutorial', 'speedrun tutorial', 'speedrun tricks', 'speedrun trick', 'mario tricks', 'mario tutorial', 'super mario 64', 
                'mexican food', 'nyc', 'pedro pascal', 'hot ones', 'mexican food nyc', 'tacos', 'new york', 'mexican street food', 'trying mexican food', 'pedro pascal interview', 'food porn', 'temu', 'tech', 'games', 'gaming', 'fake', 'switch', 'sony', 'playstation', 'shopping', 'haul', 'breath of the wild', 'tears of the kingdom', 'the legend of zelda', 'breath of the wild 2', 'legend of zelda', 'zelda tears of the kingdom', 'tears of the kingdom hype gaming', 'zelda', 'the legend of zelda tears of the kingdom', 'ocarina of time', 'the legend of zelda ocarina of time', 'tears of the kingdom theory', 'the legend of zelda: breath of the wild', 'zelda tears of the kingdom reveal', 'zelda tears of the kingdom rumors', 'wind waker', 'majoras mask', 'skyward sword', 'twilight princess', 
                'GTA', 'san andreas', 'gta 3', 'Banned', 'Commercial', 'Heineken', 'Beer', 'Jennifer', 'Aniston', 'NIRVANA', 'nirvana unplugged', 'SiriusXM', 'Sirius XM', 'Sirius', 'SXM', 'puddle of mudd', 'puddle of mudd live performance', 'Chris Van Vliet', 'CVV Clips', 'Guitar Hero', 'comedy', 'central', 'backstage stories', 'wrestling stories', 'wrestling interview', 'attitude era', 'ya fantasy books', 'ya fantasy series', 'ya fantasy book recommendations', 'ya fantasy parody', 'ya distopian', 'ya dystopian', 'Drake Bell', 'Drake and Josh', 'sonic', 'plush', 'amazon', 'sonic the hedgehog', 'game console', 'switch', 'ds', 'wii', 'iron chef', 'funny skit', 'making fun','dunkey', 'dunkey kingdom hearts', 'kingdom hearts', 'kingdom hearts explained', 'kingdom hearts recap', 'kingdom hearts story explained', 'kingdom hearts story recap', 'kingdom hearts dunkey', 
                'robot', 'live', 'concert', 'show', 'family guy', 'stewie', 'sound effect', 'tuba', 'cartoon', 'in real life', 'tall', 'short', 'son', 'mom', 'family', 'cringe compilation', 'music fails', 'musician fails', 'funny music', 'best of fails', 'top fails', 'I wrote a song', 'I wrote a song using only hate comments', 'Hate comments', 'Song', 'Songs', 'Vocals', 'Titanium', 'Comments', 'Dm', 'Direct messages', 'Comment song', 'nwa', 'goofy', 'goofy voice', 'goofy voice impression', 'disney', 'voice over', 'voiceover', 'voice actor', 'voice impressions', 'impressions', 'impression', 'WWE championship', 'World heavyweight championship', 'hidden room', 'hidden gaming room', 'secret room', 'secret gaming room', 'hidden fort', 'hidden secret fort', 'dream gaming room', 'gaming room inside stairs', 'fort under stairs', 'under stairs hidden fort', 'secret entrance inside stairs', 'Secret Room', 'Secret Rooms In Houses', 
                'Steve Austin', 'Austin 3:16', 'biker', 'bikers', 'motorcycle', 'motorcycles', 'motovlogs', 'motovloggers', 'motovlogging', 'sportbikes', 'sport bikes', 'bike life', 'arizona', 'honda', 'biker videos', 'sport bike videos', 'motorcycle blogger', 'riding a motorcycle', 'sportbike videos', 'episodes', 'series', 'movie', 'documentary', 'views', 'likes', 'subscribers', 'top', 'resident', 'evil', 'scary', 'spooky', 'epic', 'season', 'minecraft', 'fortnite', 'player', 'unknowns', 'battle', 'grounds', 'call', 'duty', 'zombies', 'awa', 'game hunting', 'video game hunting', 'retro gaming', 'retro video games', 'retro game hunting', 'best video game stores', 'video game store tour', 'retro games', 'Courage The Cowardly Dog', 'Halloween', 'Horror Game', 'Survival Horror', 'Survival Horror Game', 'deadpool', 'Marvel comics', 
                'nirvana playlist', 'nirvana clip', 'kurt cobain break string', 'kurt cobain break', 'kurt cobain scream', 'impaulsive clips', 'impaulsive', 'logan paul podcast', 'logan paul', 'maverick', 'maverick media', 'maverick house', 'logan paul clips', 'impaulsive podcast', 'hogwarts legacy', 'hogwarts legacy gameplay trailer', 'hogwarts legacy gameplay', 'harry potter game', 'harry potter', 'The Iron Claw', 'Kelly Clarkson', 'Zac Efron', 'Jeremy Allen White', 'Stanley Simons', 'Harris Dickinson', 'A24', 'he Iron Claw', 'tragic story', 'famous', 'family', 'Von Erich brothers', 'cast', 'intense training', 'transform', 'instant brothers', 'shaved', 'underwear', 'wrestle', 'The Kelly Clarkson Show', 'NBC', 'talk show', 'American Idol', 'The Voice', 'singer', 'musician', 'NBC TV', 'Television', 
                'zac efron', 'zac efron workout', 'zac efron diet', 'zac efron interview', 'zac efron body transformation', 'zac efron shirtless', 'zac efron body', 'boy and the heron', 'hayao miyazaki', 'hayao miyazaki the boy and the heron', 'new ghibli', 'new miyazaki', 'studio ghibli', 'boy meets world', 'girl meets world', 'cory matthews', 'shawn hunter', 'topanga lawrence', 'cory and shawn', 'Vince McMahon', 'McMahon', 'Jeff Hardy', 'Matt Hardy', 'Hardy Boyz', 'ichika nito', 'lofi', 'lofi guitar', 'playing god', 'ego death', 'instrumental', 'guitars', 'indie artist', 'indie music', 'gorillaz', 'gorillaz demon days', 'gorillaz cover', 'gorillaz feel good inc', 'gorillaz feel good inc cover', 'feel good inc', 'wyclef jean', 'fugees', 'lauryn hill', 'Spongebob Music', 'Spongebob Squarepants', 'Happy Spongebob', 'spin', 'spin magazine', 'spin official', 'rock', 'classic rock', 'rock music', 'hip hop', 'music industry', 
                'same voice actor', 'the same voice actor', 'same voice actor meme', 'voice actor', 'the same voice actors', 'voice actors', 'same voice actors', 'voice acting', 'same voice actor moral orel', 'they have the same voice actor', 'same voice actor regular show', 'same actor', 'voice actors everywhere', 'voice actors that are in everything', 'spider-man', 'spider-verse', 'across the spider-verse', 'spiderverse', 'gta san andreas', 'gta vice city', 'Pearl Jam', 'Imitation', 'Funny', 'Even Flow', 'Adam Sandler', 'Solo', 'Shredding', 'Hilarious', 'Epiphone', 'SG', 'School', 'Performance', 'Funniest', 'Bad', 'Worst', 'Crap', 'Worst Guitarist', 'Worst Guitar Solo Ever', 'Guitar Noob', 'patrick star', 'ai cover', 'michael jackson', 'mj', 'thriller', 'plankton', 'mr. krabs', 'frank sinatra', 'Darkness says', 'What did the five fingers', 'Slap', 'NPR', 'NPR Music', 'National Public Radio', 'Performance', 'tiny desk', 'tiny desk concert', 'tiny concert', 
                'Simpsons']

                informational_keywords = ['Spielberg', 'Steven Spielberg', 'Lipton', 'James Lipton', 'James', "Inside the Actor's Studio", 'Interview', 'Best', 'Greatest', 'Legendary', 'Question', 'Close Encounters', 'ET', 'Indiana Jones', 'Jurassic Park', 'Jaws', 'BFG', 'Favorite', 
                'Facebook engineer', 'software engineering', 'Wayne Jackson Jr', 'programming', 'software engineering journey', 'computer programmer', 'silicon valley', 'Google engineer', 'Amazon engineer', 'software development', 'software', 
                'video essay', 'screenwriting', 'analysis', 'screenplay', 'screenplays', 'screenplay tips', 'quick tips', 'how to write', 'writing', 'film', 'script', 'tips', 'story', 'review', 'filmmaking', 'advice', 'interview', 'plot', 'interviews', 'creative writing', 'script writing', 'writing tips', 'advice and tips', 'screenwriter', 'writers', 'writing process', 'masterclass', '101', 'beginners', 'motivation', 'inspiration', 'acting', 'actors', 'how to write a screenplay', 
                'vision+drive', 'visionplusdrive', 'vision plus drive', 'habits of successful people', 'how to become a millionaire', 'law of attraction', 'motivational speech', 'motivational video', 'how to get rich', 'billionaire lifestyle', 'millionaire motivation', 'motivation', 'motivational', 'success habits', 'habits of the wealthy', 'habits of the rich', 'habits of millionaires', 'success habits of millionaires', 'success mindset', 'philosophy on success', 
                'Business Insider', 'Navy SEAL', 'Navy', 'Shower', 'cold water', 'cold shower', 'video essay', 'screenwriting', 'analysis', 'screenplays', 'how to write', 'writing', 'film', 'script', 'tips', 'story', 'review', 'filmmaking', 'advice', 'interview', 'plot', 'interviews', 'creative writing', 'script writing', 'writing tips', 'advice and tips', 'screenwriter', 'writers', 'writing process', 'masterclass', '101', 'beginners', 'motivation', 'inspiration', 
                'spring boot mockito test example', 'mockito junit spring example', 'spring boot mockito controller', 'spring boot rest controller unit test example', 'mockito example', 'mockito tutorial spring boot', 'mockito service layer example', 'mockito mock dao example', 'mockito verify', 'Java Techie', 'Spring Boot', 'Mockito', 'Junit', 
                'social media', 'business', 'marketing', 'sports marketing', 'social media marekting', 'instagram', 'linkedin', 'algorithm', 'sports business', 'amazing stories', 'amazing story', 'barcroft tv', 'barcroft', 'documentary', 'real life', '2020', 'Truly', 'Johnny Quinn', 'Millington', 'Tennessee', 'burn survivor', 'fire', 'accident', 'reconstructive surgery', 'cosplay', 'third degree burns', 'body positive', 'october 2020', '246362', 'shake my beauty', 'Courageous Faces Foundation', 
                'Self Defense', 'Self Defence', 'MMA', 'Tae kwon do', 'trapping', 'training', 'fighter', 'martial artist', 'fighting', 'wrestling', 'brazillian jiu jitsu', 'jiu jitsu', 'stay safe', 'protection', 'military', 'self confidence', 'conflict management', 'personal defense', 'mixed martial arts', 'martial arts training', 'wing chun', 'muay thai', 'martial arts', 'Kung fu', 'personal development', 'boxing', 'bjj', 'karate', 'judo', 'assault', 'protect yourself', 'violence', 'empowerment', "women's self defense", 'safety', 'attack', 'Kevin goat', 'situational awareness', 'stay safe', 
                'great white shark', 'shark', 'aquarium', 'great white sharks', 'sharks', 'great white shark in aquarium', 'great white sharks in aquariums', 'great white shark in captivity', 'fish', 'animal', 'animals', 'animal shorts', 'nature', 'animal planet', 'shorts', 'communication skills', 'how to improve your communication', 'communication tips', 'how to become a better communication', 'how to improve your voice', 
                'stages of having a crush', 'crush', 'stages of love', 'having a crush', 'dating', 'love', 'relationships', 'relationship', 'crush on someone', 'psych2go', 'stage', 'fun', 'crushing', '5 stages', '6 stages of having a crush', 'psych2go crush', 'psych2go love', 'in a relationship', 'signs of a crush', 'feelings', 'emotions', 'mental health', 'psychology', 'boyfriend', 'girlfriend', 'dating advice', 'relationship advice', 'stages of relationship', 'stages of a relationship', 'romantic', 'mystery of love', 'attraction', 'romantic attraction', 'fall in love', 
                'fretless', 'fretless guitar', 'fretless guitar lesson', 'fretless guitar sound', 'how to play fretless guitar', 'prs fretless', 'PRS Guitars fretless', 'fretless prs guitar', 'Tyler Larson prs', 'prs Tyler Larson', 'music is win', 'music is win prs', 'prs music is win', 'Tyler Larson fretless', 'music is win fretless', '1 year of drumming', 'Real drum tutorial', 'know your guitar', 'know your gear', 'acoustic guitar', 'stew mac tools', 'snark tuner', 'deep dive', 'deep dive guitar review', 'phillip mcknight review', 
                'Newsflare', 'news videos', 'news today', 'journalism', 'hard news', 'worldwide', 'latest news', 'news', 'drum progress', 'progress video', 'Drum cover', 'Drum progression', 'Drum self taught', '1 year playing drums', 'Guitar progress', 'Progress video', 'Bass progress', 'Music progress', 'one year progress', 'Piano progress', 'Keyboard progress', '1 Year drums', '1 year drum', 'Drum progress 6 months', 'audio interface', 'best audio interface', 'home studio', 'best budget audio interface', 'usb audio interface', 'beginner audio interface', 
                'how rich stay rich', 'how to save money fast', 'how to save money', 'how the rich budget their money', 'how to become a millionaire', 'plumbing 101', 'plumbing basics', 'leak detection', 'leak detector', 'bath planet system', 'bathroom remodel', 'water damage restoration', 'mold removal', 'twin home experts', 'mold', 'diy', 'plumber', 'rat removal', 'rodent removal', 'rat infestation', 'twin traps', 'rat traps', 'rodent traps', 'mouse traps', 'eradicate rats', 'mouse removal', 'mouse infestation', 'twin enzyme odor removal', 'rapid mold removal', 'purocleanz', 'twinzyme', 
                'law of attraction', 'manifesting abundance', 'positive mindset', 'attracting success', 'visualization techniques', 'positive affirmations', 'attracting prosperity', 'manifestation tips', 'attracting abundance', 'attracting happiness', 'foodie', 'best snack', 'best healthy food', 'cake', 'dessert', 'healthy food', 'healthy snack for kids', 'healthy breakfast', 'the best breakfast', 'Bash', 'Bash for hackers', 'Bash for everyone', 'Bash for everybody', 'Bash right now', 'you need to learn Bash right now', 'learn Bash', 'Bash for beginners', 
                'medical animation', '3d medical animation', 'Human digestion system', 'How your body turns food into the poo', 'digestion in human beings', 'Inside the stomach', 'medical animation videos', 'digestive system', 'stomach', 'small intestine', 'rectum and anus', 'small intestine physiology', 'inside colon', 'human health', 'inside the intestine', 'digestive organs', 'human digestive system', 'poop', 'gut health', 'food digestion explained', 'large intestine', 'digestive system animation', 
                'fat loss', 'belly fat loss', 'fat loss tips', 'fat loss diet', 'how to lose belly fat', 'lose belly fat', 'fat loss workout', 'weight loss', 'how to lose fat', 'burn fat', 'belly fat', 'fat loss mistakes', 'lose fat', 'fast fat loss', 'how to lose belly fat fast', 'how to burn belly fat', 'diet for fat loss', 'fat loss journey', 'cardio for fat loss', 'burn belly fat', 'diet plan for fat loss', 'fat', 'loss', 'how to burn fat', 'weight loss tips', 'fat loss plan', 'how to lose stomach fat', 'fat loss drink', 
                'breaking', 'parallel parking', 'skill driver', 'tight parking', 'get out of tight parking spot', 'parking spot', 'how to park a car', 'parallel parking tips', 'new york tight parking', 'new york skilled driver', 'New York Parking', 'bay parking reverse', 'skill driver car', 'tight garage parking solutions', 'skilled drivers', 'parking tight spot', 'tight parking new york', 'potato recipes', 'potatoes', 'potato omelette recipe', 'potato and egg recipes', 'potato and egg', 'breakfast recipe', 'potato recipes for dinner', 'easy recipes', 'recipes', 
                'fitness', 'workouts', 'chicken recipe', 'cooking chicken', 'cooking', 'chicken', 'chicken breast', 'chicken recipes', 'asmr cooking', 'how to cook chicken', 'chicken breast recipe', 'best chicken breast', 'cooking chicken breast', 'chicken asmr cooking', 'baked chicken breast', 'cook chicken breast', 'home cooking', 'after cooking chicken', 'chicken and rice', 'chicken cooking', 'juicy chicken breast', 'how to cook chicken breast', 'Dark web', 'access the dark web', 'how to access the dark web', 'the onion network', 'the tor network', 'tor browser', 'tor network', 'tor router', 'tor website', 'edward snowden', 'tails linux', 'tails', 'tails dark web', 'dark web website', 'browsing the dark web', 
                'piano', 'piano tutorial', 'easy piano', 'piano lessons', 'piano chords', 'how to play piano', 'piano for beginners', 'piano exercises', 'piano exercises for beginners', 'beginner piano lessons', 'Luthier', 'Woodworking', 'Guitar build', 'Scrap wood', 'Hollow body', 'barncaster', 'full guitar build', 'cheap guitar', 'frugal', 'handtools', 'timelapse', 'amazon guitar build', 'amateur luthier', 'cheapest electric guitar', 'scrapwood', 'partscaster', 'scrapwood guitar', 'scrap wood guitar', 'guitar pickup', '$30', 'how to build a guitar', 'how to build guitar', 
                'Cocktail', 'recipe', 'drink', 'alcohol', 'booze', 'how to make', 'tipsy bartender', 'tipsy', 'bartender', 'liquor', 'liqueur', 'instructions', 'tutorial', 'guitar finger exercises', 'guitar exercises', 'guitar lesson', 'guitar finger exercise', 'best guitar finger exercises', 'guitar finger exercises for beginners', 'finger exercise guitar', 'finger exercises guitar', 'guitar exercise', 'guitar technique', 'economy of motion guitar', 'pinky finger guitar', 'guitar lessons', 'guitar tutorial', 'guitar finger workout', 'guitar finger training', 'guitar finger practice', 'guitar practice', 
                'street food', 'filipino food', 'asian food', 'food and travel', 'mixing vocals', 'how to mix vocals', 'Wire hangers', 'mommie dearest', 'no wire hangers', 'no wire hangers ever', 'clothes', 'clothes organization', 'Closet organization', 'declutter', 'organize', 'how to organize', 'how to organize your room', 'professional organization', 'closet', 'closet organization', 'closet organization ideas', 'small closet organization', 'stress', 'clean', 'cleaning closet', 'diy', 'small closet', 'minimalist', 'minimalist living', 'minimalist wardrobe', 'Style theory', 'fashion theory', 'theory', 
                'drum lessons', 'advanced drum lessons', 'drum fills lessons advanced', 'drum lessons for advanced drummers', 'snare drum lessons advanced', 'drum tutorial', 'drum lesson', 'drumming tutorial', 'learn drums', 'drum lessons', 'facts that could save your life', 'facts that can save your life', 'Freezer burrito recipe', 'breakfast burrito', 'breakfast meal prep', 'easy breakfast meal prep', 'freezer friendlt burritos', 'the perfect breakfast burrito', 'the best breakfast burrito', 'meal prep breakfast burritos', 'freezer meal', 'high protein burritos', 'high protein meal prep recipes', 'cheap meal prep recipes', 'body building meal prep', 
                'Presentation Skills', 'communication skills', 'business', 'training', 'success', 'coach', 'coaching', 'leadership', 'development', 'public speaking', 'speaker preparation''Presentation Skills', 'communication skills', 'business', 'training', 'success', 'coach', 'coaching', 'leadership', 'development', 'public speaking', 'speaker preparation', 'Rats', 'Tickling', 'Giggles', 'Neuroscience', 'science', 'Nvidia', 'Graphics Card', 'GPU', '3090 Ti', 'Titanium', 'Ti', 'Pronunciation', 'Word', 'How To', 'Gaming', 'PC', 'Desktop', 'Hardware', 
                'Blender', 'Mantaflow', 'Flip Fluids', 'Green Man', 'Fluid', 'Simulation', 'Realistic', 'HD', 'High Res', 'Particles', 'Mario', 'Glass', 'Pipe', 'New', 'HDR', 'HDRi', 'Cycles', 'Mesh', 'Shade Smooth', 'Smooth', 'Ray Tracing', 'Ray', 'RTX', 'Particle', 'Satisfying', 'Render', 'Animation', 'Reflection', 'Water Simulation', 'Fluid Simulation', '60fps', 'refinance', 'mortgage', 'mortgages', 'refinancing', 'mortgage rates', 'refinance mortgage', 'refinance rates', 'refinancing rates', 'refinancing home', 'home loan', 'home loans', 'equity loans', 'home equity loans', 'home equity loan', 'second mortgage', 'home equity loan rates', 'credit card consolidation', 'debt loans', 'credit card debt consolidation', 'bad credit loans', 'debt free', 'bad credit', 'debt help', 'debt solutions', 'money management', 'credit card debt', 'personal loan', 'bad credit mortgage', 'mortgage calculator', 
                'music theory for guitarists', 'guitar music theory', 'guitar theory', 'beginner guitar theory', 'beginner music theory', 'music theory lesson guitar', 'beginner music theory guitar', 'theory guitar', 'music theory guitar tutorial', 'music theory tricks', 'music theory lessons', 'music theory tutorial', 'music theory tutorials', 'music theory exercises guitar', 'music theory lessons guitar', 'music theory lessons for guitarists', 'intermediate music theory', 'advanced music theory', 'guitar tutorial', 'guitar lesson', 'guitar technique', 'guitar exercises', 'guitar strings', 'beginner guitarist', 'pc', 'pc tips', 'pc building', 'samsung', '4k monitor', 'odyssey', 'neo g9', 'dual 4k monitor', 'gaming monitor', 'review', 
                'donuts recipe', 'donut recipe', 'easy donuts', 'easy donuts recipe', 'no yeast donuts', 'donuts no yeast', 'donuts recipe no yeast', 'eggless donuts', 'homemade donuts recipe', 'homemade donuts', 'how to make donuts', 'glazed donuts', 'glazed donuts recipe', 'bread recipe', 'hygiene routines', 'hygiene tips', 'air fryer donuts', 'air fryer donuts recipe', 'air fryer donuts from scratch', 'air fryers donuts easy', 'the best air fryer donuts', 'glazed donuts', 'glazed donuts recipe', 'how to make donuts in an airfryer', 'krispy kreme donuts recipe', 'glazed donuts from scratch', 
                'Kitchen Knife', 'sharpest kitchen knife', 'sharpest knife', 'razor sharp', 'razor blade', 'utility blade', 'homemade knife', 'knife', 'replicable blades', 'great idea', 'sharpening', 'best knife', 'make a knife', 'forging', 'worlds sharpest', 'sharpner', 'knife making', 'coolest knife', 'worlds sharpest kitchen knife', 'sharpest in the world', 'sharpest knife in the world', 'sharpest kitchen knife in the world', 'picking guitar', 'picking guitar lesson', 'guitar picking exercises', 'picking exercises', 'alternate picking', 'alternate picking lesson', 'alternate picking exercises', 'alternate picking tutorial', 'economy picking', 'sweep picking', 'shred guitar lesson', 'shred', 'learn', 'learning', 'hack', 
                'waterjet channel', 'water jet', 'water jet cutting', 'waterjet', 'water cutting', 'water cutter', 'cutting with water', 'cutting', 'cut', 'cross section', 'cross-section', 'cut in half', 'inside', 'interesting', 'science', 'fascinating', 'satisfying videos', 'amazing inventions', 'sex', 'relationship', 'premature ejaculation', 'orgasm', 'better sex coach', 'erectile dysfunction', 'ED', 'good sex', 'tv show good sex', 'bodybuilding', 'fitness', 'build muscle', 'gym', 'workout', 'hypertrophy', 'muscle building', 'muscle growth', 'fat loss', 'burn fat', 'weight loss', 'fat burning', 'diet', 'nutrition', 
                'best paid jobs', 'business', 'management', 'financial management', 'project management', 'ceo', 'personal finance', 'career planning', 'small business', 'why are managers bad at their jobs', 'business management', 'top ceo', 'worst ceo', 'micromanager', 'bad manager', 'job creator', 'bull jobs', 'working', 'career', 'easiest jobs', 'hardest jobs', 'finance', 'easiest jobs to get', '40 hour work week', 'wage vs salary', 'work from home', 'jobs that do nothing', 'unemployment', 'boring job', 'easy job', 'easiest job for pay', 'passion job', 'anxious', 'self aware', 'stress', 'self help', 'psychology', 'therapy', 'health', 'mental health', 
                'how it\'s actually made', 'how it\'s made', 'cooking video', 'cooking demo', 'kitchen hack', 'cooking tip', 'easy recipe', 'quick recipe', 'food science', 'brickbending', 'brick', 'bending', 'LEGO', 'legoart', 'geometricart', 'brick bending', 'lego bending', 'my own creation', 'legos', 'moc', 'afol', 'lego', 'how to', 'lego build', 'geometry', 'diy', 'meditative', 'satisfying lego', 'oddly satisfying', 'legobending', 'bricks', 'diy beautiful', 'creative', 'amazing lego', 'satisfying', 'engineering', 'lego brick', 'sculpture', 'creativity', 'howto', 'lego bricks', 'geometric art', 'lego art', 'lego moc', 'lego sculpture', 'asmr', 
                'how peanut butter is made', 'cooking', 'food', 'chicken', 'fried chicken', 'orange chicken', 'easy recipe', 'auto', 'auto repair', 'car', 'car diy', 'car repair', 'cars', 'how to', 'mechanic', 'repair', 'car review', 'car theft', 'car thief', 'how cars are stolen', 'how to keep car from being stolen', 'most stolen', 'most stolen cars', 'motor vehicle theft', 'stolen', 'stolen car', 'stolen cars', 'stolen cars caught on camera', 'stolen cars found', 'education', 'philosophy', 'talk', 'self', 'improvement', 'big questions', 'love', 'wellness', 'mindfullness', 'psychology', 'how', 'to', 'hack', 'relationship', 'relationship advice', 'relationship repair', 'couples', 'communication', 'trust', 'intimacy', 'honesty', 'self-improvement', 'breakup', 'divorce', 'infidelity', 'jealousy', 'anger', 'conflict', 
                'knife sharpening', 'pull through knife shaprener', 'knife edge close up', 'knife sharpening easy', 'carbide knife sharpener', 'sharpens best', 'knife shaprening close up', 'sharpening', 'knife sharpening up close', 'close up knife edge', 'close up sharpening', 'pull through knife sharpener', 'how to use a pull through knife sharpener', 'extrovert meets introvert', 'introvert', 'introverts', '16 personalities', 'introvert humor', 'introvert memes', 'introvert problems', 'extrovert', 'introvert vs extrovert', 'songwriting', 'songwriting tips', 'songwriting process', 'songwriting techniques', 'clear ketchup', 'clear recipes', 'transparent ketchup', 'clear tomato ketchup', 'transparent tomato ketchup', 'tomato ketchup', 'transparent recipes', 'see through recipes', 'how to make clear ketchup', 'homemade clear ketchup', 
                'Chocolate Cake', 'Cake', 'easy recipes', 'chocolate cake recipe', 'cake recipe', 'cake decorating', 'snickers recipe', 'how to make snickers', 'snickers at home', 'candy bar recipe', 'how to make candy', 'homemade candy', 'dessert recipes', 'caramel recipe', 'how to make caramel', 'homemade caramel', 'caramel sauce', 'dark chocolate recipes', 'homemade snickers bar', 'edible science', 'cookery channel', 'my virgin kitchen', 'clear foods', 'how to make transparent bacon', 'how to make bacon', 'how to', 'bacon salt', 'bacon stock', 'vegan bacon', 'clear bacon', 'transparent bacon', 'glass bacon', 'see through bacon', 'see through recipes', 'clear recipes', 'lonely', 'signs of loneliness', 'dr julie', 'tasting history', 'food history', 'typescript', 'type script', 'typescript type predicate', 'type predicate', 'type predicate function', 'ts', 'ts type predicate', 'ts type predicate function', 'typescript type predicate function', 
                'Art Support', 'Dave Grohl interview', 'Dave Grohl advice', 'the power of gratitude', 'Dave Grohl motivational speech', 'foo fighters', 'how to be a badass', 'How to Use The Inspiration to Become a Badass Inspirer', 'How to Unlock Your True Potential Like a Badass', 'motivational speech', 'Career advice', 'crayola colored pencils', 'blending colored pencils', 'how to blend colored pencils', 'blending cheap colored pencils', 'how to make cheap colored pencils look smooth', 'how to use crayola colored pencils', 'how to make cheap colored pencils look expensive', 'how to make cheap colored pencils better', 'apartment', 'tiniest apartment', 'tiny apartment tokyo lens', 'tiny apartment Japan', 'tiny apartment tokyo', 'tiny apartment', 'small apartment', 'micro apartment', 'small apartment japan', "japan's smallest apartment", "japan's tiniest apartment", 'weird apartment', 'worst apartment', 'japans worst apartment', 'small apartment tour', 'weird', 'weird japan', 'triangle apartment', 'narrow apartment', 'apartment tour', 
                'bluetooth', 'how does bluetooth work', 'ihop pancakes recipe', 'homemade pancake recipe', 'homemade pancake mix', 'pancake mix', 'ihop pancakes', 'buttermilk pancake recipe', 'perfect pancake recipe', 'homemade pancakes', 'pancakes', 'how to make pancakes', 'buttermilk pancakes', 'fluffy pancakes', 'easy pancakes', 'how to make the best pancakes at home', 'cooking', 'korean food', 'asmr cooking', 'cooking video', 'cooking vlog', 'easy recipe', 'korean', 'korean culture', 'foodie', 'mukbang', 'emotional beat making', 'making a beat from scratch', 'beat making from scratch', 'making a beat', 'how to land a tech job', 'how to land a developer job', 'developer job', 'how to land a dev job', 'javascript job', 'web dev job', 'web developer job', 'how to become a web developer', 'how to get hired', 
                'learn to code', 'get a job', 'learn programming', 'programming', 'coding', 'how to code', 'software developer', 'software engineering', 'software developer job', 'How to learn', 'How to learn how to code', 'Get a job in tech', 'tech jobs with no experience', 'tech jobs', 'software engineer', 'code', 'coding job', 'javascript', 'learn javascript', 'how to program an app', 'startup', 'computer science', 'learn how to code', 'internship interview', 'tech job', 'learn to code for beginners', 'no cs degree', 'self taught programmer', 'how to learn to code', 'how to learn programming', 'things to never say in an interview', 'job interview mistakes', 'interview preparation', 'job interview questions', 'job interview tips', 'job interview questions and answers', 'interview tips', 'interview questions and answers', 'job interview', 'what to say in a job interview', 'hiring manager', 'interview help', 'what not to say in an interview', 'how to prepare for an interview', 'interview questions', 'interview blunders', 'how to pass a job interview', 'interview mistakes', 
                'Public Speaking', 'Public Speaking Tips', 'Public Speaking Training', 'Public Speaking Class', 'Public Speaking Anxiety', 'Public Speaking Techniques', 'Public Speaking Course', 'Public Speaking Examples', 'Presentation Skills', 'How to Give a Speech', 'Business Presentation', 'Persuasive Speech Example', 'College Speech', 'react 19', 'react changes', 'react 19 changes', 'reactjs 19', 'react19 js', 'react js', 'reactjs', 'react', 'react javascript', 'react js 19', 'react javascript 19', 'react next version', 'new react hooks', 'new react features', 'camping', 'camping tips', 'canoe', 'canoe bailer', 'canoe bailing kit', 'canoe flip', 'canoe instruction', 'canoe rescue', 'canoe rescue technique', 'canoe rescue techniques', 'canoe rescues', 'canoe self rescue', 'canoe tutorial', 'knots', 'camping knots', 'best survival knots', 'knot', 'knots you need to know', 'bushcraft', 'survival', 'hitch', 'how to tie', 'bowline', 'square knot', 'prepping', 'life hacks', 'timber hitch', 'camping tips', 'how to tie a knot', 'climbing knots', 'basic knots', 'essential knots', 'beginner knots', 'outdoor knots', 'knots for camping', 'alpine butterfly', 'zeppelin bend', '69 knot', 'sheet bend', 'taut line', 'adjustable', 'prusik', 'evenk hitch', 'truckers hitch', 'jam knot', 'constrictor knot', 'survival skills', 'useful knots', 
                'adhd', 'hyperactive', 'nervousness', 'palpitations', 'adrenaline', 'cortisol', 'insomnia', 'worried', 'tension', 'wound up', 'fearful', 'difficulty sleeping', 'headaches', 'neck pain', 'tinnitus', 'vertigo', 'pinched nerve', 'herniated disc', 'hand reflexology', 'acupressure', 'acupuncture', 'stuffy nose', 'back pain', 'sinusitus', 'sinus', 'clogged nose', 'migraine', 'brain', 'nerves', 'eyes', 'ringing in ears', 'nervous', 'depression', 'how to destress', 'sleep quickly', 'how to relax', 'reduce stress', 'calmness', 'peace', 'sleep', 'how to get rich', 'make money online', 'how to become a millionaire', 'how to make money online', 'online business', 'how to make money', 'how to build wealth', 'how to be a millionaire', 'entrepreneur motivation', 'small business ideas', 'business ideas', 'chicken tenders recipe', 'chicken recipe', 'fried chicken tenders', 'chicken fingers recipe', 'how to make chicken tenders', 'how to cook chicken tenders', 'best chicken tender recipe', 'how to cook', 'crispy fried chicken recipe', 'homemade ranch', 'ranch', 'homemade ranch seasoning', 'homemade ranch salad dressing', 'homemade ranch dressing', 
                'how its made', 'how it\'s made', 'how do they do it', 'how to make', 'factory made', 'how it\'s made full episodes', 'how it\'s made episodes', 'everyday items', 'production line', 'mechanical production', 'discovery', 'how stuff works', 'how its made episodes', 'discovery channel', 'how its made food', 'making of', 'made', 'cpu', 'intel', 'amd', 'things i wish i knew', 'life advice', 'life lessons', 'makeup', 'make-up']

                
                r = list_paragraph2.add_run()

                r.add_text(title_content)
                list_paragraph2.add_run('\n')
                r.add_picture(image_parsed, width=Pt(300))
                list_paragraph2.add_run('\n')
                hyperlink = add_hyperlink(list_paragraph2, url, url, 'FF8822', False)
                list_paragraph2.add_run('\n')
                list_paragraph2.add_run('\n')
                list_paragraph2.add_run('\n')
                list_paragraph2.add_run('\n')
                list_paragraph2.add_run('\n')

                print(title_content + " informational")
                

                """
                    elif word in entertainment_keywords:
                        r = list_paragraph1.add_run()

                        r.add_text(title_content)
                        list_paragraph1.add_run('\n')
                        r.add_picture(io_url, width=Pt(300))
                        list_paragraph1.add_run('\n')
                        hyperlink = add_hyperlink(list_paragraph1, url, url, 'FF8822', False)
                        list_paragraph1.add_run('\n')
                        list_paragraph1.add_run('\n')
                        list_paragraph1.add_run('\n')
                        list_paragraph1.add_run('\n')
                        list_paragraph1.add_run('\n')

                        print(title_content + " entertainment")
                        break

                    elif word in general_keywords:
                        r = list_paragraph3.add_run()

                        r.add_text(title_content)
                        list_paragraph3.add_run('\n')
                        r.add_picture(io_url, width=Pt(300))
                        list_paragraph3.add_run('\n')
                        hyperlink = add_hyperlink(list_paragraph3, url, url, 'FF8822', False)
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')

                        print(title_content + " general")
                        break
                
                    else:
                        r = list_paragraph3.add_run()

                        r.add_text(title_content)
                        list_paragraph3.add_run('\n')
                        r.add_picture(io_url, width=Pt(300))
                        list_paragraph3.add_run('\n')
                        hyperlink = add_hyperlink(list_paragraph3, url, url, 'FF8822', False)
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')
                        list_paragraph3.add_run('\n')

                        print(title_content + " general")
                """

                


                """
                all_keys = set()
                key_holder = []
                if keywords:
                    for keyword in keywords:
                        keys = keyword.split(', ')
                        for key in keys:
                            key_str = str(key)
                            key_holder.append(key_str)
                print(key_holder)
                """
                
            else:
                print("Failed to retrieve the web page. Status code: {response.status_code}")
            
        except:
            print("An error occurred")

        doc1.save('entertainment.docx')
        doc2.save("informational.docx")
        doc3.save("general.docx")

        

# Be sure to close the connection
con.close()